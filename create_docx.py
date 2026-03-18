"""RAG 학습 자료 DOCX 생성 스크립트"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 스타일 설정 ──
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# 한글 폰트 설정
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), '맑은 고딕')
rPr.append(rFonts)

for level, size, color in [
    ('Heading 1', 22, '1F4E79'),
    ('Heading 2', 16, '2E75B6'),
    ('Heading 3', 13, '2E75B6'),
]:
    s = doc.styles[level]
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = True
    s.font.name = 'Arial'
    rPr = s.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.append(rFonts)
    s.paragraph_format.space_before = Pt(18)
    s.paragraph_format.space_after = Pt(8)


def add_code_block(text):
    """코드 블록 스타일 단락 추가"""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # 배경색 (연한 회색)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        run.element.get_or_add_rPr().append(shd)


def add_tip_box(title, text):
    """팁/참고 박스"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    # 왼쪽 테두리
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '2E75B6')
    pBdr.append(left)
    pPr.append(pBdr)
    run_title = p.add_run(f'{title}\n')
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    run_title.font.size = Pt(11)
    run_body = p.add_run(text)
    run_body.font.size = Pt(10)


def add_table(headers, rows):
    """표 추가"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()  # 간격


def add_bullet(text, level=0):
    """불릿 리스트"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.space_after = Pt(2)


def add_numbered(text, level=0):
    """번호 리스트"""
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.space_after = Pt(2)


# ================================================================
#  표지
# ================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('RAG 완전 정복 가이드')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Retrieval-Augmented Generation')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('초보자를 위한 단계별 학습 자료')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
lines = [
    '개념 이해부터 실전 구현까지',
    '임베딩, 벡터DB, LangChain 활용법',
    '한국어 예제와 실습 코드 포함',
]
for line in lines:
    run = p.add_run(line + '\n')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ================================================================
#  목차
# ================================================================
doc.add_heading('목차', level=1)
toc_items = [
    ('Chapter 1.  RAG란 무엇인가?', '개념, 필요성, 전체 흐름'),
    ('Chapter 2.  임베딩과 벡터DB 이해하기', '임베딩, 코사인 유사도, 텍스트 분할'),
    ('Chapter 3.  단계별 실습 튜토리얼', '환경 설정부터 완성까지'),
    ('Chapter 4.  용어 사전', 'RAG 필수 용어 정리'),
    ('Chapter 5.  자주 하는 실수와 트러블슈팅', '8가지 흔한 실수와 해결법'),
    ('Chapter 6.  학습 로드맵', '다음 단계 가이드'),
]
for title, desc in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(f'\n    {desc}').font.size = Pt(10)

doc.add_page_break()

# ================================================================
#  Chapter 1: RAG란 무엇인가?
# ================================================================
doc.add_heading('Chapter 1. RAG란 무엇인가?', level=1)

add_tip_box(
    '한 줄 요약',
    'RAG(Retrieval-Augmented Generation)는 LLM이 답변할 때, '
    '외부 문서에서 관련 정보를 검색(Retrieval)해서 함께 참고하도록 하는 기술입니다.'
)

doc.add_heading('왜 RAG가 필요한가?', level=2)
doc.add_heading('LLM의 한계', level=3)

add_bullet('학습 데이터 이후의 정보를 모른다 (최신 뉴스, 업데이트 등)')
add_bullet('우리 회사 내부 문서를 모른다 (사내 정책, 제품 매뉴얼)')
add_bullet('없는 정보를 그럴듯하게 지어낸다 (할루시네이션)')
add_bullet('출처를 알 수 없다 (어디서 나온 정보인지 확인 불가)')

doc.add_heading('RAG가 해결하는 것', level=3)

add_bullet('최신 문서를 검색해서 답변 -> 정보가 항상 최신')
add_bullet('우리 회사 문서를 기반으로 답변 -> 내부 지식 활용')
add_bullet('문서에 근거한 답변만 생성 -> 할루시네이션 감소')
add_bullet('참조한 문서를 함께 보여줌 -> 출처 확인 가능')

doc.add_heading('RAG vs 파인튜닝 vs 프롬프트 엔지니어링', level=2)

add_table(
    ['구분', '프롬프트 엔지니어링', 'RAG', '파인튜닝'],
    [
        ['방법', '질문할 때 정보를 직접 넣어줌', '벡터DB에서 검색해서 넣어줌', '모델 자체를 재학습'],
        ['비유', '시험 볼 때 커닝페이퍼', '시험 볼 때 교과서 검색', '교과서를 외우게 함'],
        ['비용', '거의 없음', '중간 (임베딩 비용)', '높음 (GPU 비용)'],
        ['최신성', '수동 업데이트', '문서만 업데이트하면 됨', '재학습 필요'],
        ['적합한 경우', '정보가 적을 때', '대량의 문서 기반 QA', '모델 행동을 바꿀 때'],
    ]
)

add_tip_box('초보자 추천', '대부분의 경우 RAG로 시작하는 것이 가장 효율적입니다.')

doc.add_heading('RAG의 전체 흐름 (Big Picture)', level=2)

doc.add_paragraph('RAG는 크게 2단계로 나뉩니다:')

doc.add_heading('1단계: 인덱싱 (Indexing) - 사전 준비', level=3)
doc.add_paragraph(
    '문서를 미리 처리해서 검색 가능한 형태로 저장합니다.'
)

add_code_block(
    '[원본 문서]          [텍스트 청크]           [벡터(숫자)]          [벡터 DB]\n'
    '   PDF    --Load-->  "연차는..."  --Embed-->  [0.1, 0.3, ...]  --Store-->  ChromaDB\n'
    '   TXT               "재택은..."             [0.2, 0.1, ...]               FAISS\n'
    '   Web                "급여는..."             [0.5, 0.4, ...]               Pinecone'
)

doc.add_heading('2단계: 질의응답 (Query) - 실제 사용', level=3)
doc.add_paragraph(
    '사용자가 질문하면 관련 문서를 찾아서 LLM에 함께 전달합니다.'
)

add_code_block(
    '[사용자 질문]                      [관련 문서 검색]                [LLM 답변 생성]\n'
    '"연차 며칠?"  --Embed--> 벡터  --Search--> "연차는 15일..."  --LLM--> "연차는 15일입니다"\n'
    '                         |                  "3년 이상시..."             + 출처 표시\n'
    '                    [0.1, 0.3, ...]'
)

doc.add_heading('비유로 이해하기: RAG = 똑똑한 비서 + 도서관', level=2)

p = doc.add_paragraph()
run = p.add_run('기존 LLM:\n')
run.bold = True
doc.add_paragraph('  학생: "우리 회사 연차 정책이 뭐야?"')
doc.add_paragraph('  LLM:  "일반적으로 회사 연차는..." (일반적인 답변, 부정확)')

p = doc.add_paragraph()
run = p.add_run('RAG 적용 후:\n')
run.bold = True
doc.add_paragraph('  학생: "우리 회사 연차 정책이 뭐야?"')
doc.add_paragraph('  비서:  (도서관에서 \'연차\' 관련 문서를 빠르게 찾음)')
doc.add_paragraph('  비서:  (찾은 문서를 LLM에게 건네줌)')
doc.add_paragraph('  LLM:  "우리 회사 규정에 따르면, 1년 이상 근속 시 연 15일입니다"')

doc.add_heading('핵심 구성 요소', level=2)

add_table(
    ['구성 요소', '역할', '예시'],
    [
        ['Document Loader', '문서를 읽어오는 도구', 'TextLoader, PyPDFLoader'],
        ['Text Splitter', '긴 문서를 작은 조각으로 분할', 'RecursiveCharacterTextSplitter'],
        ['Embedding Model', '텍스트를 숫자 벡터로 변환', 'OpenAI text-embedding-3-small'],
        ['Vector Store', '벡터를 저장/검색하는 DB', 'ChromaDB, FAISS, Pinecone'],
        ['Retriever', '질문과 유사한 문서를 검색', 'vectorstore.as_retriever()'],
        ['LLM', '검색 결과로 답변을 생성', 'GPT-4o-mini, Claude'],
    ]
)

doc.add_page_break()

# ================================================================
#  Chapter 2: 임베딩과 벡터DB 이해하기
# ================================================================
doc.add_heading('Chapter 2. 임베딩과 벡터DB 이해하기', level=1)

add_tip_box(
    '이 챕터의 목표',
    '임베딩이 뭔지, 왜 필요한지, 벡터DB가 어떻게 "의미 기반 검색"을 가능하게 하는지 이해합니다.'
)

doc.add_heading('1. 임베딩(Embedding)이란?', level=2)

doc.add_paragraph(
    '텍스트를 고차원 숫자 벡터로 변환하는 것입니다. '
    '의미가 비슷한 텍스트는 비슷한 벡터가 됩니다.'
)

doc.add_heading('왜 숫자로 바꿔야 하나?', level=3)
doc.add_paragraph(
    '컴퓨터는 텍스트의 "의미"를 직접 이해하지 못합니다. '
    '하지만 숫자라면 비교할 수 있습니다.'
)

add_code_block(
    '"고양이는 귀엽다"  -->  [0.82, 0.15, 0.43, 0.91, ...]  --+\n'
    '                                                          |-- 숫자가 비슷! = 의미가 비슷!\n'
    '"냥이는 사랑스럽다" -->  [0.80, 0.17, 0.45, 0.89, ...]  --+\n'
    '\n'
    '"오늘 날씨가 좋다"  -->  [0.12, 0.67, 0.33, 0.05, ...]  <-- 숫자가 다름 = 의미가 다름'
)

doc.add_heading('임베딩의 핵심 특성', level=3)

add_numbered('고차원 벡터: 텍스트 하나가 수백~수천 개의 숫자로 변환됨 (예: 1536차원)')
add_numbered('의미적 유사성 보존: "사과"와 "배"의 벡터는 가까움 (둘 다 과일)')
add_numbered('방향성: 왕 - 남자 + 여자 = 여왕 (벡터 연산으로 의미 관계 표현 가능)')

doc.add_heading('비유: 도서관의 분류 번호', level=3)

add_table(
    ['검색 방식', '동작', '결과'],
    [
        ['키워드 검색', '책 제목에 "고양이" 글자가 있는 책 찾기', '"냥이", "cat" 검색 안 됨'],
        ['임베딩 검색', '고양이에 "관한" 책 찾기', '"냥이", "cat", "반려묘" 모두 찾음'],
    ]
)

doc.add_heading('2. 코사인 유사도 (Cosine Similarity)', level=2)

doc.add_paragraph('두 벡터가 얼마나 비슷한지 측정하는 방법입니다.')

add_table(
    ['값', '의미', '해석'],
    [
        ['1.0', '완전히 같은 방향', '의미가 거의 같음'],
        ['0.0', '직각', '관계 없음'],
        ['-1.0', '반대 방향', '의미가 반대'],
    ]
)

doc.add_heading('Python으로 확인해보기', level=3)

add_code_block(
    'from langchain_openai import OpenAIEmbeddings\n'
    'import numpy as np\n'
    '\n'
    'embeddings = OpenAIEmbeddings(model="text-embedding-3-small")\n'
    '\n'
    'texts = ["고양이는 귀엽다", "냥이는 사랑스럽다", "오늘 날씨가 좋다"]\n'
    'vectors = embeddings.embed_documents(texts)\n'
    '\n'
    'def cosine_similarity(a, b):\n'
    '    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))\n'
    '\n'
    'print(f"고양이 vs 냥이: {cosine_similarity(vectors[0], vectors[1]):.4f}")  # 높음!\n'
    'print(f"고양이 vs 날씨: {cosine_similarity(vectors[0], vectors[2]):.4f}")  # 낮음!'
)

doc.add_heading('3. 텍스트 분할 (Text Splitting)', level=2)

doc.add_heading('왜 분할하나?', level=3)

add_bullet('문서가 너무 길면 임베딩 품질이 떨어짐 (의미가 희석됨)')
add_bullet('LLM 컨텍스트 윈도우에 제한이 있음')
add_bullet('적절한 크기로 분할하면 검색 정확도가 높아짐')

doc.add_heading('분할 과정 시각화', level=3)

add_code_block(
    '원본 문서 (1000자)\n'
    '  "1. 근무시간 정책 ... 2. 휴가 정책 ... 3. 복리후생 ..."\n'
    '       |\n'
    '       v  RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=50)\n'
    '       |\n'
    '  +----------+  +----------+  +----------+\n'
    '  | 청크 1    |  | 청크 2    |  | 청크 3    |\n'
    '  | 근무시간  |  | 휴가 정책 |  | 복리후생   |\n'
    '  | (300자)   |  | (300자)   |  | (250자)   |\n'
    '  +----------+  +----------+  +----------+\n'
    '         ^^^^----^^^^  <-- overlap(겹침): 문맥 끊김 방지'
)

doc.add_heading('분할 파라미터 가이드', level=3)

add_table(
    ['파라미터', '작게 설정', '크게 설정', '추천 시작값'],
    [
        ['chunk_size', '정밀한 검색, 문맥 부족', '풍부한 문맥, 노이즈 증가', '300~500'],
        ['chunk_overlap', '문맥 끊김 위험', '중복 증가, 저장공간 증가', 'chunk_size의 10~20%'],
    ]
)

doc.add_heading('4. 벡터 DB (Vector Database)', level=2)

doc.add_paragraph(
    '벡터(숫자 배열)를 저장하고, "가장 비슷한 벡터"를 빠르게 검색할 수 있는 데이터베이스입니다.'
)

doc.add_heading('일반 DB vs 벡터 DB', level=3)

add_table(
    ['구분', '일반 DB (MySQL)', '벡터 DB (Chroma)'],
    [
        ['검색 방식', '정확히 일치하는 것만 검색', '의미가 비슷한 것 검색'],
        ['쿼리 예시', 'WHERE title = "연차"', 'search("연차가 궁금해요", k=3)'],
        ['"휴가" 검색 가능?', '불가능 (키워드 불일치)', '가능 (의미 유사)'],
    ]
)

doc.add_heading('주요 벡터 DB 비교', level=3)

add_table(
    ['벡터 DB', '특징', '적합한 경우'],
    [
        ['Chroma', '설치 쉬움, 로컬 파일 저장', '학습, 프로토타입, 소규모'],
        ['FAISS', 'Meta 제작, 매우 빠름', '대규모 검색, 로컬'],
        ['Pinecone', '클라우드 서비스, 관리형', '프로덕션, 서버리스'],
        ['Weaviate', '오픈소스, 풍부한 기능', '중대규모 프로덕션'],
        ['Qdrant', 'Rust 기반, 고성능', '고성능 프로덕션'],
    ]
)

add_tip_box('초보자 추천', 'Chroma로 시작하고, 익숙해지면 FAISS나 Pinecone으로 전환하세요.')

doc.add_heading('5. 검색 전략', level=2)

add_table(
    ['전략', '설명', '장점', '단점'],
    [
        ['Similarity', '코사인 유사도 기반 (기본)', '단순하고 빠름', '비슷한 내용 중복 가능'],
        ['MMR', '관련성 + 다양성 고려', '중복 없이 다양한 정보', '약간 느림'],
        ['Score Threshold', '점수 기준으로 필터링', '관련 없는 문서 제외', '임계값 조정 필요'],
    ]
)

doc.add_page_break()

# ================================================================
#  Chapter 3: 단계별 실습 튜토리얼
# ================================================================
doc.add_heading('Chapter 3. 단계별 실습 튜토리얼', level=1)

add_tip_box(
    '이 챕터의 목표',
    '코드를 한 줄씩 따라 치면서 RAG를 직접 구현합니다. 각 단계마다 "왜 이렇게 하는지"를 함께 설명합니다.'
)

doc.add_heading('사전 준비', level=2)

doc.add_heading('1) 환경 설정', level=3)

add_code_block(
    '# 가상환경 생성 (권장)\n'
    'python -m venv venv\n'
    'source venv/bin/activate       # Mac/Linux\n'
    'venv\\Scripts\\activate          # Windows\n'
    '\n'
    '# 패키지 설치\n'
    'pip install langchain langchain-openai langchain-community chromadb tiktoken python-dotenv'
)

doc.add_heading('2) API 키 설정', level=3)

add_code_block(
    '# .env 파일 생성\n'
    'OPENAI_API_KEY=sk-your-actual-api-key'
)

doc.add_paragraph('API 키가 없다면 https://platform.openai.com/api-keys 에서 발급받으세요.')

doc.add_heading('실습 1: 임베딩 체험하기', level=2)
doc.add_paragraph('목표: 텍스트가 어떻게 숫자(벡터)로 변환되는지 직접 확인합니다.')

add_code_block(
    'from dotenv import load_dotenv\n'
    'from langchain_openai import OpenAIEmbeddings\n'
    'import numpy as np\n'
    '\n'
    'load_dotenv()\n'
    '\n'
    '# 임베딩 모델 생성\n'
    'embeddings = OpenAIEmbeddings(model="text-embedding-3-small")\n'
    '\n'
    '# 하나의 텍스트를 벡터로 변환\n'
    'text = "고양이는 귀여운 동물이다"\n'
    'vector = embeddings.embed_query(text)\n'
    '\n'
    'print(f"원본 텍스트: {text}")\n'
    'print(f"벡터 차원 수: {len(vector)}")   # 1536\n'
    'print(f"벡터 앞 5개: {vector[:5]}")\n'
    '\n'
    '# 여러 텍스트의 유사도 비교\n'
    'sentences = [\n'
    '    "고양이는 귀여운 동물이다",       # 원본\n'
    '    "냥이는 사랑스러운 반려동물이다",  # 의미 유사\n'
    '    "강아지도 귀여운 동물이다",        # 부분 유사\n'
    '    "오늘 주식이 폭락했다",            # 관련 없음\n'
    ']\n'
    'vectors = embeddings.embed_documents(sentences)\n'
    '\n'
    'def cosine_sim(a, b):\n'
    '    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))\n'
    '\n'
    'for i in range(1, len(sentences)):\n'
    '    sim = cosine_sim(vectors[0], vectors[i])\n'
    '    print(f"  {sim:.4f}  {sentences[i]}")'
)

p = doc.add_paragraph()
run = p.add_run('배운 것: ')
run.bold = True
p.add_run('embed_query()는 검색 질문용(1개), embed_documents()는 문서 저장용(여러 개). '
           '의미가 비슷하면 유사도가 높다.')

doc.add_heading('실습 2: 문서 로드 & 분할', level=2)
doc.add_paragraph('목표: 실제 문서를 읽고 적절한 크기로 분할합니다.')

add_code_block(
    'from langchain_community.document_loaders import TextLoader\n'
    'from langchain.text_splitter import RecursiveCharacterTextSplitter\n'
    '\n'
    '# 문서 로드\n'
    'loader = TextLoader("sample_documents/company_policy.txt", encoding="utf-8")\n'
    'documents = loader.load()\n'
    '\n'
    'print(f"문서 수: {len(documents)}")\n'
    'print(f"전체 길이: {len(documents[0].page_content)}자")\n'
    '\n'
    '# 텍스트 분할\n'
    'splitter = RecursiveCharacterTextSplitter(\n'
    '    chunk_size=200,\n'
    '    chunk_overlap=30,\n'
    '    separators=["\\n\\n", "\\n", ". ", " ", ""],\n'
    ')\n'
    'chunks = splitter.split_documents(documents)\n'
    '\n'
    'print(f"분할 결과: {len(chunks)}개 청크")\n'
    'for i, chunk in enumerate(chunks[:3]):\n'
    '    print(f"  [청크 {i+1}] ({len(chunk.page_content)}자)")\n'
    '    print(f"  {chunk.page_content[:60]}...")'
)

p = doc.add_paragraph()
run = p.add_run('배운 것: ')
run.bold = True
p.add_run('TextLoader로 Document 객체로 변환, RecursiveCharacterTextSplitter가 가장 범용적, '
           'separators 순서대로 자연스러운 위치에서 분할.')

doc.add_heading('실습 3: 벡터 저장소 만들기', level=2)
doc.add_paragraph('목표: 분할된 청크를 벡터DB에 저장하고 검색합니다.')

add_code_block(
    'from langchain_openai import OpenAIEmbeddings\n'
    'from langchain_community.vectorstores import Chroma\n'
    '\n'
    '# 벡터 저장소 생성 (임베딩 + 저장을 한 번에!)\n'
    'embeddings = OpenAIEmbeddings(model="text-embedding-3-small")\n'
    'vectorstore = Chroma.from_documents(\n'
    '    documents=chunks,\n'
    '    embedding=embeddings,\n'
    ')\n'
    '\n'
    '# 검색 테스트\n'
    'query = "연차 휴가"\n'
    'results = vectorstore.similarity_search(query, k=2)\n'
    'for i, doc in enumerate(results):\n'
    '    print(f"  결과{i+1}: {doc.page_content[:80]}...")\n'
    '\n'
    '# 유사도 점수와 함께 검색\n'
    'results_with_score = vectorstore.similarity_search_with_score(query, k=3)\n'
    'for doc, score in results_with_score:\n'
    '    print(f"  거리: {score:.4f} | {doc.page_content[:60]}...")'
)

p = doc.add_paragraph()
run = p.add_run('배운 것: ')
run.bold = True
p.add_run('Chroma.from_documents()로 한 번에 임베딩+저장, '
           'similarity_search()로 유사 문서 검색, '
           'similarity_search_with_score()로 점수 확인 가능.')

doc.add_heading('실습 4: 완성! 질문-답변 RAG', level=2)
doc.add_paragraph('목표: 검색 결과를 LLM에 전달하여 자연어 답변을 생성합니다.')

add_code_block(
    'from langchain_openai import ChatOpenAI\n'
    'from langchain.chains import RetrievalQA\n'
    'from langchain.prompts import PromptTemplate\n'
    '\n'
    '# 프롬프트 설계\n'
    'prompt = PromptTemplate(\n'
    '    input_variables=["context", "question"],\n'
    '    template="""당신은 회사 정책을 안내하는 AI 어시스턴트입니다.\n'
    '아래 [참고 문서]만을 근거로 질문에 답변하세요.\n'
    '문서에 없는 내용은 "해당 정보는 확인되지 않습니다"라고 답하세요.\n'
    '\n'
    '[참고 문서]\n'
    '{context}\n'
    '\n'
    '[질문] {question}\n'
    '[답변]""",\n'
    ')\n'
    '\n'
    '# RAG 체인 조립\n'
    'llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)\n'
    'qa_chain = RetrievalQA.from_chain_type(\n'
    '    llm=llm,\n'
    '    chain_type="stuff",       # 검색 결과를 통째로 넣기\n'
    '    retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),\n'
    '    return_source_documents=True,\n'
    '    chain_type_kwargs={"prompt": prompt},\n'
    ')\n'
    '\n'
    '# 질문하기!\n'
    'result = qa_chain.invoke({"query": "신입사원의 연차는 어떻게 되나요?"})\n'
    'print(f"A: {result[\'result\']}")\n'
    'print(f"참조 문서: {len(result[\'source_documents\'])}개")'
)

p = doc.add_paragraph()
run = p.add_run('배운 것: ')
run.bold = True
p.add_run('RetrievalQA 체인으로 검색+답변을 하나로 묶음, '
           'chain_type="stuff"는 가장 단순한 방식, '
           '커스텀 프롬프트로 답변 형식과 톤을 제어, '
           '문서에 없는 질문에 "모른다"고 답하게 유도 가능.')

doc.add_heading('실습 5: 대화형 RAG (챗봇)', level=2)
doc.add_paragraph('목표: "아까 말한 거" 같은 이전 대화를 기억하는 RAG 챗봇을 만듭니다.')

add_code_block(
    'from langchain.memory import ConversationBufferMemory\n'
    'from langchain.chains import ConversationalRetrievalChain\n'
    '\n'
    '# 대화 메모리\n'
    'memory = ConversationBufferMemory(\n'
    '    memory_key="chat_history",\n'
    '    return_messages=True,\n'
    '    output_key="answer",\n'
    ')\n'
    '\n'
    '# 대화형 RAG 체인\n'
    'chain = ConversationalRetrievalChain.from_llm(\n'
    '    llm=ChatOpenAI(model="gpt-4o-mini", temperature=0),\n'
    '    retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),\n'
    '    memory=memory,\n'
    '    return_source_documents=True,\n'
    ')\n'
    '\n'
    '# 연속 대화 (맥락을 기억함!)\n'
    'chain.invoke({"question": "SmartWidget Pro 요금제를 알려주세요"})\n'
    'chain.invoke({"question": "그 중에서 AI 분석이 포함된 건?"})  # "그 중에서" = 이전 맥락!\n'
    'chain.invoke({"question": "그 플랜의 가격은?"})               # "그 플랜" = 이전 답변!'
)

p = doc.add_paragraph()
run = p.add_run('배운 것: ')
run.bold = True
p.add_run('ConversationBufferMemory로 대화 기록 저장, '
           'ConversationalRetrievalChain으로 대화 맥락 + 문서 검색 결합, '
           '"그 중에서", "아까 말한" 같은 참조를 이해.')

doc.add_page_break()

# ================================================================
#  Chapter 4: 용어 사전
# ================================================================
doc.add_heading('Chapter 4. 용어 사전', level=1)
doc.add_paragraph('RAG를 공부하면서 만나는 용어들을 쉽게 정리했습니다.')

terms = [
    ('기본 용어', [
        ('LLM', 'Large Language Model. ChatGPT, Claude 같은 대규모 언어 모델.'),
        ('RAG', 'Retrieval-Augmented Generation. LLM이 답변 전에 관련 문서를 검색해서 참고하는 기술.'),
        ('할루시네이션', 'LLM이 사실이 아닌 내용을 그럴듯하게 지어내는 현상. RAG로 감소 가능.'),
    ]),
    ('임베딩 관련', [
        ('임베딩 (Embedding)', '텍스트를 고차원 숫자 벡터로 변환하는 것. 의미 유사 = 벡터 유사.'),
        ('벡터 (Vector)', '숫자의 배열. [0.1, 0.3, 0.5, ...] 형태. 임베딩의 결과물.'),
        ('코사인 유사도', '두 벡터의 유사도 측정 (-1~1). 1에 가까울수록 의미가 유사.'),
        ('임베딩 모델', '텍스트를 벡터로 변환하는 AI 모델. 예: text-embedding-3-small'),
    ]),
    ('문서 처리 관련', [
        ('Document Loader', '다양한 형식의 파일을 Document 객체로 변환하는 도구.'),
        ('Document', 'page_content(내용) + metadata(출처 등)으로 구성된 문서 단위.'),
        ('청크 (Chunk)', '긴 문서를 분할한 작은 텍스트 조각.'),
        ('Text Splitter', '문서를 청크로 분할하는 도구. RecursiveCharacterTextSplitter 추천.'),
        ('청크 오버랩', '인접 청크 간 겹치는 부분. 문맥 끊김 방지용.'),
    ]),
    ('벡터 저장소 관련', [
        ('벡터 DB', '벡터를 저장하고 유사도 기반 검색을 제공하는 데이터베이스.'),
        ('인덱스 (Index)', '벡터를 효율적으로 검색할 수 있도록 정리한 구조.'),
        ('메타데이터', '문서에 붙이는 부가 정보. 카테고리, 날짜 등으로 필터링 가능.'),
    ]),
    ('검색 관련', [
        ('리트리버 (Retriever)', '질문과 관련된 문서를 검색하는 도구.'),
        ('Top-K', '검색 결과에서 상위 K개만 반환. k=3이면 가장 유사한 3개.'),
        ('MMR', 'Maximal Marginal Relevance. 관련성 높고 다양한 내용의 문서 선택.'),
        ('시맨틱 검색', '의미 기반 검색. 키워드가 아니라 뜻이 통하는 문서를 찾음.'),
    ]),
    ('LLM 체인 관련', [
        ('체인 (Chain)', '여러 작업을 순서대로 연결한 파이프라인.'),
        ('프롬프트 템플릿', 'LLM에 보낼 프롬프트의 틀. {변수}로 동적 내용 삽입.'),
        ('chain_type', '검색 결과를 LLM에 전달하는 방식. stuff(통째로), map_reduce(개별처리) 등.'),
        ('Temperature', 'LLM 답변의 창의성 조절 (0~2). RAG에는 0 추천.'),
    ]),
    ('대화형 RAG 관련', [
        ('대화 메모리', '이전 대화 내용을 저장하여 맥락을 유지하는 기능.'),
        ('대화 히스토리', '지금까지 주고받은 대화 기록. "그것", "아까" 등을 이해하게 해줌.'),
    ]),
]

for category, items in terms:
    doc.add_heading(category, level=2)
    for term, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(f'{term}: ')
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        p.add_run(desc)

doc.add_page_break()

# ================================================================
#  Chapter 5: 자주 하는 실수와 트러블슈팅
# ================================================================
doc.add_heading('Chapter 5. 자주 하는 실수와 트러블슈팅', level=1)
doc.add_paragraph('RAG를 처음 구현할 때 누구나 겪는 문제들과 해결 방법을 정리했습니다.')

mistakes = [
    (
        '실수 1: chunk_size가 너무 크거나 작음',
        '검색 결과가 엉뚱하거나, 답변이 너무 일반적이거나, 중요한 세부 정보가 누락됨.',
        '300~500자로 시작하고 결과를 보면서 조정. chunk_overlap을 chunk_size의 10~20%로 설정. '
        '분할 결과를 직접 눈으로 확인하여 각 청크가 완결된 의미를 담고 있는지 점검.'
    ),
    (
        '실수 2: 검색된 문서가 엉뚱함',
        '"연차 정책"을 물었는데 "제품 요금제"가 검색되거나, 관련 없는 문서가 상위에 올라옴.',
        'k 값을 2~4로 줄여보기. 메타데이터로 카테고리 분류 후 필터링 검색. '
        'similarity_search_with_score()로 점수를 확인하여 디버깅.'
    ),
    (
        '실수 3: API 키 에러',
        'openai.AuthenticationError: Incorrect API key provided',
        '.env 파일 존재 확인, 따옴표/공백 없이 키 입력, load_dotenv() 호출 확인, '
        'API 키 유효성 확인 (https://platform.openai.com/api-keys).'
    ),
    (
        '실수 4: 한국어 인코딩 에러',
        'UnicodeDecodeError 발생',
        'TextLoader에 encoding="utf-8" 지정. Windows 파일이면 encoding="cp949" 시도. '
        'BOM 포함 파일이면 encoding="utf-8-sig" 사용.'
    ),
    (
        '실수 5: 답변이 문서 내용을 무시함',
        '문서에 "연차 15일"이라고 써있는데 LLM이 일반적인 지식으로 답변.',
        '프롬프트에 "반드시 참고 문서만을 근거로 답변하세요", '
        '"문서에 없는 내용은 확인되지 않습니다라고 답하세요"를 명시적으로 지시.'
    ),
    (
        '실수 6: 메모리(RAM) 부족',
        '프로그램이 느려지거나 MemoryError 발생.',
        '문서를 배치로 나눠서 처리. persist_directory로 디스크 저장 활용. '
        '대규모 데이터는 FAISS나 Pinecone 고려. 벡터DB를 재사용하여 임베딩 비용 절약.'
    ),
    (
        '실수 7: PDF 텍스트 추출 실패',
        'PDF를 로드했는데 빈 텍스트가 나오거나 텍스트가 깨짐.',
        '스캔된 이미지 PDF는 OCR 라이브러리 사용 (pytesseract). '
        'PyPDFLoader 대신 PDFMinerLoader나 UnstructuredPDFLoader 시도.'
    ),
    (
        '실수 8: 비용이 너무 많이 나옴',
        '임베딩과 LLM API 호출 비용이 예상보다 많음.',
        'text-embedding-3-small 사용 (large 대비 5배 저렴), '
        'gpt-4o-mini 사용 (gpt-4o 대비 30배 저렴), '
        '벡터DB를 persist해서 임베딩 재생성 방지.'
    ),
]

for title, symptom, solution in mistakes:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    run = p.add_run('증상: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    p.add_run(symptom)

    p = doc.add_paragraph()
    run = p.add_run('해결: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
    p.add_run(solution)

doc.add_heading('디버깅 체크리스트', level=2)
doc.add_paragraph('문제가 생겼을 때 순서대로 확인하세요:')

checklist = [
    '.env 파일에 API 키가 올바르게 들어있는가?',
    'load_dotenv()를 호출했는가?',
    '문서가 정상적으로 로드되었는가? (documents 출력해서 확인)',
    '청크가 적절한 크기로 분할되었는가? (chunks 출력해서 확인)',
    '검색 결과가 질문과 관련 있는가? (similarity_search 직접 확인)',
    '프롬프트가 문서 참조를 강제하고 있는가?',
    'LLM temperature가 0인가? (사실 기반 답변에 추천)',
]
for item in checklist:
    add_numbered(item)

doc.add_page_break()

# ================================================================
#  Chapter 6: 학습 로드맵
# ================================================================
doc.add_heading('Chapter 6. 학습 로드맵', level=1)

doc.add_heading('단계별 학습 경로', level=2)

levels = [
    ('Level 1: 기초 이해 (1~2일)', [
        'RAG 개념과 필요성 이해',
        '임베딩과 벡터 개념 이해',
        '기본 RAG 파이프라인 따라하기',
    ]),
    ('Level 2: 기본 구현 (2~3일)', [
        '문서 로드 & 분할 실습',
        '벡터DB 저장 & 검색 실습',
        '전체 RAG 파이프라인 조립',
        '프롬프트 커스터마이징',
    ]),
    ('Level 3: 심화 기능 (3~5일)', [
        '대화형 RAG (챗봇)',
        '고급 검색 전략 (MMR, 멀티쿼리)',
        '메타데이터 필터링',
        'PDF 등 다양한 문서 형식 처리',
    ]),
    ('Level 4: 실전 프로젝트 (1~2주)', [
        '자신의 문서로 RAG 시스템 구축',
        'Streamlit/Gradio로 웹 UI 만들기',
        '평가 지표로 품질 측정',
        '프로덕션 고려사항 (비용, 성능, 보안)',
    ]),
    ('Level 5: 고급 아키텍처 (지속 학습)', [
        'LCEL (LangChain Expression Language)',
        'Agent + RAG 결합',
        'Hybrid Search (키워드 + 시맨틱)',
        'Re-ranking, HyDE, Graph RAG',
    ]),
]

for level_title, items in levels:
    doc.add_heading(level_title, level=3)
    for item in items:
        add_bullet(item)

doc.add_heading('프로젝트 아이디어', level=2)

doc.add_heading('초급', level=3)
add_numbered('내 노트 검색기 - 자신의 메모/노트 파일로 RAG 구축')
add_numbered('FAQ 챗봇 - 자주 묻는 질문 목록으로 자동 답변 봇')
add_numbered('독서 노트 도우미 - 읽은 책의 핵심 내용 저장 및 검색')

doc.add_heading('중급', level=3)
add_numbered('기술 문서 검색 시스템 - API 문서, 기술 블로그 검색 + 웹 UI')
add_numbered('법률/규정 QA 시스템 - 사내 규정 기반 QA + 메타데이터 필터링')
add_numbered('다국어 RAG - 영문 문서를 한국어로 질문하여 답변')

doc.add_heading('고급', level=3)
add_numbered('멀티모달 RAG - 이미지, 표, 차트 포함 문서 처리')
add_numbered('Agent + RAG - RAG를 도구로 사용하는 AI 에이전트')
add_numbered('실시간 RAG - 웹 크롤링 + 벡터DB 자동 업데이트 파이프라인')

doc.add_heading('핵심 라이브러리 정리', level=2)

add_table(
    ['라이브러리', '역할', '설치 명령어'],
    [
        ['langchain', 'RAG 프레임워크', 'pip install langchain'],
        ['langchain-openai', 'OpenAI 연동', 'pip install langchain-openai'],
        ['langchain-community', '커뮤니티 통합', 'pip install langchain-community'],
        ['chromadb', '벡터 DB', 'pip install chromadb'],
        ['faiss-cpu', '벡터 DB (고속)', 'pip install faiss-cpu'],
        ['pypdf', 'PDF 처리', 'pip install pypdf'],
        ['tiktoken', '토큰 카운팅', 'pip install tiktoken'],
        ['streamlit', '웹 UI', 'pip install streamlit'],
    ]
)

doc.add_heading('학습 팁', level=2)

tips = [
    '코드를 복사하지 말고 직접 타이핑하세요 - 손으로 치면서 이해가 깊어집니다.',
    '파라미터를 바꿔가며 실험하세요 - chunk_size, k, temperature 등을 바꿔보면 동작 원리를 체감합니다.',
    '자신의 문서로 시도하세요 - 샘플 문서 대신 실제 관심 있는 문서로 테스트하세요.',
    '에러를 두려워하지 마세요 - 에러 메시지는 최고의 학습 자료입니다.',
    '중간 결과를 항상 출력해보세요 - print()로 각 단계의 결과를 확인하면 흐름이 보입니다.',
]
for tip in tips:
    add_numbered(tip)

# ── 저장 ──
output_path = 'C:/Users/ilove/dev/rag-sample/RAG_학습_가이드.docx'
doc.save(output_path)
print(f"DOCX 파일 생성 완료: {output_path}")
